"""
Agent Conversation Renderer — Streamlit 版（OpenAI）
AG-UI Protocol：模型回傳結構化 JSON，在聊天泡泡內渲染對應元件
"""

import os
import re
import json
import uuid
import base64
import hashlib
import textwrap
from pathlib import Path
from datetime import datetime

import streamlit as st
from streamlit_ace import st_ace
import openai

# ══════════════════════════════════════════════════════════════
# 頁面設定
# ══════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Agent Conversation Renderer",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed",
)

UPLOAD_DIR    = Path("uploads")
GENERATED_DIR = Path("generated")
UPLOAD_DIR.mkdir(exist_ok=True)
GENERATED_DIR.mkdir(exist_ok=True)

DEFAULT_MODEL = st.secrets.get("OPENAI_MODEL", "gpt-3.5")

# ══════════════════════════════════════════════════════════════
# Session state
# ══════════════════════════════════════════════════════════════
def init_state():
    defaults = {
        "messages":        [],   # [{role, content?, agui?, file_chip?}]
        "uploaded_files":  [],
        "generated_files": [],
        "active_file":     None,
        "thinking":        False,
        "client":          None,
        "pending_suggestion": "",
        "thinking_stages": [],
        "chat_expanded":   True,   # 新增：對話欄展開狀態
        "stage_expanded":  False,  # 新增：展示欄展開狀態（初始折疊）
        "split_ratio":     25,     # 新增：對話欄占比（25% = 折疊展示欄時的初始狀態）
        "api_key_override": "",    # 新增：使用者輸入的 API Key
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# ══════════════════════════════════════════════════════════════
# OpenAI client
# ══════════════════════════════════════════════════════════════
def get_client() -> openai.OpenAI:
    if st.session_state.client is None:
        api_key  = st.secrets.get("OPENAI_API_KEY", "")
        base_url = st.secrets.get("OPENAI_BASE_URL", None)
        if not api_key:
            st.error("請在 `.streamlit/secrets.toml` 設定 `OPENAI_API_KEY`")
            st.stop()
        st.session_state.client = openai.OpenAI(api_key=api_key, base_url=base_url)
    return st.session_state.client

# ══════════════════════════════════════════════════════════════
# 檔案工具
# ══════════════════════════════════════════════════════════════
def file_icon(ext: str) -> str:
    return {"pdf":"📄","docx":"📝","doc":"📝","png":"🖼️","jpg":"🖼️",
            "jpeg":"🖼️","gif":"🖼️","webp":"🖼️","svg":"🖼️"}.get(ext.lower(),"📎")

def is_image(ext): return ext.lower() in {"png","jpg","jpeg","gif","webp"}
def is_pdf(ext):   return ext.lower() == "pdf"
def is_docx(ext):  return ext.lower() in {"docx","doc"}

def file_b64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode()

def save_upload(up) -> dict:
    fid  = str(uuid.uuid4())[:8]
    ext  = Path(up.name).suffix.lstrip(".")
    dest = UPLOAD_DIR / f"{fid}_{up.name}"
    dest.write_bytes(up.getvalue())
    orig_content = ""
    if ext.lower() in {"html", "md", "svg", "json", "txt"}:
        try:
            orig_content = up.getvalue().decode("utf-8", errors="replace")
        except:
            pass
    return {"id":fid,"name":up.name,"path":dest,"ext":ext,
            "size":up.size,"ts":datetime.now().strftime("%H:%M"),"source":"upload",
            "original_content": orig_content}

def save_generated(name, content_bytes, ext, prompt="") -> dict:
    fid  = str(uuid.uuid4())[:8]
    dest = GENERATED_DIR / f"{fid}_{name}"
    dest.write_bytes(content_bytes)
    orig_content = ""
    try:
        orig_content = content_bytes.decode("utf-8", errors="replace")
    except:
        pass
    return {"id":fid,"name":name,"path":dest,"ext":ext,
            "ts":datetime.now().strftime("%H:%M"),"source":"model","prompt":prompt,
            "original_content": orig_content}

def find_file(fid) -> dict | None:
    return next((f for f in st.session_state.uploaded_files +
                 st.session_state.generated_files if f["id"] == fid), None)

def is_editable_text(ext: str) -> bool:
    return ext.lower() in {"html", "md", "svg", "json", "txt"}

def generate_diff_html(original: str, current: str) -> str:
    import difflib

    orig_lines = original.splitlines()
    curr_lines = current.splitlines()

    unified = list(difflib.unified_diff(
        orig_lines, curr_lines,
        fromfile="原始版本 (Original)",
        tofile="目前版本 (Current)",
        lineterm="",
        n=3,
    ))

    def escape(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    rows_html = []
    line_no_old = 0
    line_no_new = 0

    if not unified:
        rows_html.append(
            '<tr><td class="ln"></td><td class="ln"></td>'
            '<td class="ctx">（無差異，兩版本完全相同）</td></tr>'
        )
    else:
        for line in unified:
            if line.startswith("---") or line.startswith("+++"):
                rows_html.append(
                    f'<tr><td class="ln"></td><td class="ln"></td>'
                    f'<td class="hdr">{escape(line)}</td></tr>'
                )
            elif line.startswith("@@"):
                # 解析 hunk header 取得行號
                import re
                m = re.search(r"-(\d+)(?:,\d+)? \+(\d+)", line)
                if m:
                    line_no_old = int(m.group(1)) - 1
                    line_no_new = int(m.group(2)) - 1
                rows_html.append(
                    f'<tr><td class="ln"></td><td class="ln"></td>'
                    f'<td class="hunk">{escape(line)}</td></tr>'
                )
            elif line.startswith("-"):
                line_no_old += 1
                rows_html.append(
                    f'<tr><td class="ln del">{line_no_old}</td><td class="ln"></td>'
                    f'<td class="del">− {escape(line[1:])}</td></tr>'
                )
            elif line.startswith("+"):
                line_no_new += 1
                rows_html.append(
                    f'<tr><td class="ln"></td><td class="ln add">{line_no_new}</td>'
                    f'<td class="add">＋ {escape(line[1:])}</td></tr>'
                )
            else:
                line_no_old += 1
                line_no_new += 1
                rows_html.append(
                    f'<tr><td class="ln">{line_no_old}</td><td class="ln">{line_no_new}</td>'
                    f'<td class="ctx">　 {escape(line[1:] if line.startswith(" ") else line)}</td></tr>'
                )

    rows = "\n".join(rows_html)

    html = f"""<!DOCTYPE html>
<html lang="zh-Hant">
<head>
<meta charset="UTF-8">
<style>
  html, body {{
    margin: 0;
    padding: 0;
    background: #faf8f4;
    color: #1d2a2a;
    font-family: Consolas, "Noto Sans Mono", monospace;
    font-size: 13px;
    overflow-x: hidden;
    overflow-y: auto;
  }}

  .diff-wrap {{
    padding: 8px;
  }}

  table {{
    width: 100%;
    border-collapse: collapse;
    table-layout: fixed;
  }}

  /* 行號欄 */
  td.ln {{
    width: 38px;
    min-width: 38px;
    max-width: 38px;
    text-align: right;
    padding: 2px 6px 2px 4px;
    color: #aaa;
    border-right: 1px solid #e8e0d4;
    user-select: none;
    vertical-align: top;
    font-size: 11px;
    line-height: 1.6;
  }}

  /* 內容欄：自動換行，不橫向溢出 */
  td:last-child {{
    padding: 2px 10px;
    white-space: pre-wrap;
    word-break: break-word;
    overflow-wrap: anywhere;
    vertical-align: top;
    line-height: 1.6;
  }}

  /* 各行類型配色 */
  tr.del-row td.ln   {{ background: #fff0f0; color: #c0392b; }}
  tr.del-row td:last-child {{ background: #fff0f0; color: #c0392b; }}

  tr.add-row td.ln   {{ background: #f0fff4; color: #27ae60; }}
  tr.add-row td:last-child {{ background: #f0fff4; color: #27ae60; }}

  tr.ctx-row td:last-child {{ background: #faf8f4; color: #555; }}

  tr.hunk-row td:last-child {{
    background: #eef4ff;
    color: #3b5bdb;
    font-size: 11px;
    padding: 4px 10px;
  }}

  tr.hdr-row td:last-child {{
    background: #f0ebe1;
    color: #888;
    font-size: 11px;
    padding: 4px 10px;
    border-bottom: 1px solid #e8e0d4;
  }}

  /* 行間隔線 */
  tr {{ border-bottom: 1px solid rgba(216, 205, 189, 0.25); }}

  /* 圖例 */
  .legend {{
    display: flex;
    gap: 16px;
    padding: 6px 10px;
    background: #f0ebe1;
    border-radius: 8px;
    margin-bottom: 8px;
    font-size: 11px;
    color: #666;
    flex-wrap: wrap;
  }}
  .legend span {{ display: flex; align-items: center; gap: 4px; }}
  .dot {{ width: 10px; height: 10px; border-radius: 3px; display: inline-block; }}
  .dot-del {{ background: #fde8e8; border: 1px solid #c0392b; }}
  .dot-add {{ background: #d4f0d4; border: 1px solid #27ae60; }}
  .dot-ctx {{ background: #faf8f4; border: 1px solid #ccc; }}
</style>
</head>
<body>
<div class="diff-wrap">
  <div class="legend">
    <span><span class="dot dot-del"></span> 刪除</span>
    <span><span class="dot dot-add"></span> 新增</span>
    <span><span class="dot dot-ctx"></span> 未變更</span>
  </div>
  <table>
    <colgroup>
      <col style="width:38px">
      <col style="width:38px">
      <col>
    </colgroup>
    <tbody>
{rows}
    </tbody>
  </table>
</div>
</body>
</html>"""

    # 把 class 名稱注入 tr（因為是字串拼接，直接替換 class 標記）
    html = html.replace('<td class="del">', '<td class="del">')  # no-op, classes already on td
    # 修正 tr class
    import re as _re
    def _fix_tr(m):
        inner = m.group(1)
        if 'class="del"' in inner or 'class="ln del"' in inner:
            return f'<tr class="del-row">{inner}</tr>'
        elif 'class="add"' in inner or 'class="ln add"' in inner:
            return f'<tr class="add-row">{inner}</tr>'
        elif 'class="hunk"' in inner:
            return f'<tr class="hunk-row">{inner}</tr>'
        elif 'class="hdr"' in inner:
            return f'<tr class="hdr-row">{inner}</tr>'
        else:
            return f'<tr class="ctx-row">{inner}</tr>'

    html = _re.sub(r'<tr>(.*?)</tr>', _fix_tr, html, flags=_re.DOTALL)
    return html

def detect_artifact_kind(text: str, language: str = "") -> tuple[str, str] | None:
    stripped = text.strip()
    lang = language.lower().strip()

    if not stripped:
        return None
    if lang == "html" or stripped.lower().startswith("<!doctype html") or stripped.lower().startswith("<html"):
        return "html", stripped
    if lang == "svg" or stripped.lower().startswith("<svg"):
        return "svg", stripped
    if lang in {"json", "application/json"}:
        return "json", stripped
    if lang in {"markdown", "md"}:
        return "md", stripped
    if stripped.startswith("{") and stripped.endswith("}"):
        try:
            json.loads(stripped)
            return "json", stripped
        except json.JSONDecodeError:
            return None
    return None

def make_artifact_name(ext: str, title: str = "") -> str:
    stem = re.sub(r"[^A-Za-z0-9_-]+", "_", title).strip("_") or f"artifact_{ext}"
    return f"{stem}.{ext}"

def upsert_generated_artifact(name: str, content: str, ext: str, prompt: str = "") -> dict:
    content_bytes = content.encode("utf-8")
    digest = hashlib.md5(content_bytes).hexdigest()
    existing = next(
        (
            f for f in st.session_state.generated_files
            if f["ext"] == ext and f.get("content_hash") == digest
        ),
        None,
    )
    if existing:
        st.session_state.active_file = existing["id"]
        return existing

    meta = save_generated(name, content_bytes, ext, prompt=prompt)
    meta["content_hash"] = digest
    st.session_state.generated_files.append(meta)
    st.session_state.active_file = meta["id"]
    return meta

# ══════════════════════════════════════════════════════════════
# AG-UI System Prompt
# ══════════════════════════════════════════════════════════════
AGUI_SYSTEM = textwrap.dedent("""
你是一個 AI agent，在雙欄式介面運作。回答時必須回傳一個 JSON，格式嚴格符合 AGUIResponse。

## AGUIResponse 格式
```json
{
  "components": [ ...一或多個 component... ],
  "suggestions": ["下一個問題1", "下一個問題2"]
}
```

## 可用 component 類型（從中選最適合的）

### markdown
```json
{"type":"markdown","content":"Markdown 格式文字"}
```
用於：一般說明、解釋、段落文字。

### info_card
```json
{"type":"info_card","title":"標題","description":"說明文字","variant":"info"}
```
variant 可選：info | warning | success | danger
用於：重要提示、警告、狀態通知。

### data_list
```json
{"type":"data_list","title":"標題（選填）","items":[{"label":"欄位","value":"內容"}]}
```
用於：key-value 對、屬性清單、詳細資訊。

### step_process
```json
{"type":"step_process","title":"流程標題","steps":[{"title":"步驟1","description":"說明"}]}
```
用於：操作步驟、流程說明、教學。

### table
```json
{"type":"table","title":"表格標題","headers":["欄1","欄2"],"rows":[["A","B"]]}
```
用於：比較資料、結構化清單、多欄資訊。

### stat_grid
```json
{"type":"stat_grid","title":"統計標題","items":[{"label":"指標","value":"數值","description":"說明（選填）"}]}
```
用於：KPI、數據摘要、指標展示。

### code_block
```json
{"type":"code_block","title":"標題（選填）","language":"python","content":"程式碼內容"}
```
用於：程式碼、指令、設定檔。

### action_group
```json
{"type":"action_group","title":"動作標題","items":[{"label":"動作名稱","action":"action_id","description":"說明（選填）"}]}
```
用於：建議操作、快速動作清單。

### surface
```json
{"type":"surface","kind":"html","html":"<html>...</html>","css":"body{}","title":"標題"}
```
kind 可選：html | svg | markdown | iframe
用於：生成完整 HTML 應用介面、SVG 圖表、需要展示在右側舞台的大型內容。
若包含 surface component，它會自動顯示在右側展示介面。

## 規則
1. 【最重要】你的回應必須是且只能是一個合法的 JSON 物件，從 `{` 開始，到 `}` 結束。絕對禁止在 JSON 前後加任何文字、說明、markdown 包裝或 ``` 符號。
2. 即使使用者的輸入不完整、缺乏上下文或無法執行，你仍然必須回傳 JSON，使用 `info_card`（variant: warning）說明問題，並在 `suggestions` 引導使用者提供更多資訊。
3. 根據問題內容自動選擇最適合的 component 組合
4. 可以在同一個 components 陣列內混用多種類型
5. 繁體中文回答
6. suggestions 提供 2-3 個後續可問的問題
""").strip()

# ══════════════════════════════════════════════════════════════
# 三階段 OpenAI 調用函數
# ══════════════════════════════════════════════════════════════

def call_openai_stage1_text_content(user_text: str, attached_file: dict | None = None) -> dict:
    """Stage 1: 生成純文本內容和建議"""
    client = get_client()
    model  = st.secrets.get("OPENAI_MODEL", DEFAULT_MODEL)

    system_prompt = textwrap.dedent("""
    你是一個 AI 助手，以繁體中文回答。

    你的任務：
    1. 清楚、準確地回答使用者的問題
    2. 提供 0-2 個相關的後續問題建議

    回傳一個 JSON 物件（只包含 answer 和 suggestions，不包含 components）：
    {
      "answer": "你的回答文字",
      "suggestions": ["建議問題1", "建議問題2"]
    }

    【重要】你的回應必須是合法的 JSON，從 { 開始到 } 結束，無其他文字。
    """).strip()

    history: list[dict] = [{"role": "system", "content": system_prompt}]

    # 構建訊息內容（支持檔案處理）
    msg_content: list = []

    if attached_file:
        p   = attached_file["path"]
        ext = attached_file["ext"].lower()

        if is_image(ext):
            mime = f"image/{'jpeg' if ext in ('jpg','jpeg') else ext}"
            msg_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{file_b64(p)}"},
            })
            msg_content.append({"type": "text", "text": user_text or "請分析這張圖片。"})
        elif is_pdf(ext):
            txt = pdf_extract_text(p)
            msg_content.append({"type": "text", "text": f"[PDF 內容]\n{txt[:4000]}\n\n{user_text}"})
        elif is_docx(ext):
            txt = docx_extract_text(p)
            msg_content.append({"type": "text", "text": f"[DOCX 內容]\n{txt[:4000]}\n\n{user_text}"})
        elif ext in {"txt", "md"}:
            try:
                file_text = Path(p).read_text(encoding="utf-8", errors="replace")
            except Exception as e:
                file_text = f"（檔案讀取失敗：{e}）"
            fname = attached_file.get("name", f"file.{ext}")
            msg_content.append({
                "type": "text",
                "text": f"[{fname} 內容]\n{file_text[:8000]}\n\n{user_text}",
            })
        else:
            fname = attached_file.get("name", f"file.{ext}")
            msg_content.append({
                "type": "text",
                "text": f"（使用者上傳了檔案：{fname}，格式暫不支援解析）\n\n{user_text}",
            })
    else:
        msg_content.append({"type": "text", "text": user_text})

    history.append({"role": "user", "content": msg_content})

    # API call
    resp = client.chat.completions.create(model=model, messages=history, temperature=0.7, max_tokens=2048)
    raw  = resp.choices[0].message.content

    # 清除 ``` 包裝
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```\s*$", "", raw, flags=re.MULTILINE)
    raw = raw.strip()

    # 擷取第一個 { 到最後一個 }
    start = raw.find("{")
    end   = raw.rfind("}")
    if start != -1 and end != -1 and end > start:
        raw = raw[start:end + 1]

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {
            "answer": raw[:500] or "（無法生成回答）",
            "suggestions": [],
        }


def call_openai_stage2_component_plan(answer_text: str) -> dict:
    """Stage 2: 決定使用哪些 components"""
    client = get_client()
    model  = st.secrets.get("OPENAI_MODEL", DEFAULT_MODEL)

    component_types = "markdown, info_card, data_list, step_process, table, stat_grid, code_block, action_group"

    system_prompt = textwrap.dedent(f"""
    你是一個 UI 元件選擇專家。

    根據提供的回答文本，決定使用哪些 component 類型最合適。
    目標是讓回答以視覺化、結構化方式呈現，不要把大部分內容都放進 markdown。

    可用的 component 類型：
    - markdown：一般說明、解釋、段落文字、Mermaid 圖表
    - info_card：重要提示、警告、狀態通知
    - data_list：key-value 對、屬性清單、詳細資訊
    - step_process：操作步驟、流程說明、教學
    - table：比較資料、結構化清單、多欄資訊
    - stat_grid：KPI、數據摘要、指標展示
    - code_block：程式碼、指令、設定檔
    - action_group：建議操作、快速動作清單

    回傳 JSON：
    {{
      "components_to_use": ["info_card", "data_list", "code_block"],
      "component_descriptions": {{
        "info_card": "摘要或關鍵提醒",
        "data_list": "結構化重點",
        "code_block": "程式碼範例"
      }}
    }}

    選型規則：
    1. 優先選 2 到 4 種互補 component。
    2. markdown 只用於前言、結語、短段落補充，不可作為唯一主體，除非內容真的無法結構化。
    3. 如果是教學、操作、流程，優先包含 step_process。
    4. 如果是比較、差異、清單整理，優先包含 table 或 data_list。
    5. 如果有數字、指標、等級、狀態摘要，優先包含 stat_grid 或 info_card。
    6. 如果有程式碼、指令、設定片段，優先包含 code_block。
    7. 如果內容包含明確建議行動，可包含 action_group。

    【重要】回應必須是合法的 JSON，從 {{ 開始到 }} 結束。
    """).strip()

    history = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"回答文本：\n{answer_text}"},
    ]

    resp = client.chat.completions.create(model=model, messages=history, temperature=0.5, max_tokens=1024)
    raw  = resp.choices[0].message.content

    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```\s*$", "", raw, flags=re.MULTILINE)
    raw = raw.strip()

    start = raw.find("{")
    end   = raw.rfind("}")
    if start != -1 and end != -1 and end > start:
        raw = raw[start:end + 1]

    try:
        result = json.loads(raw)
        # 驗證 components_to_use 是有效的類型
        valid_types = {c.lower() for c in component_types.split(", ")}
        result["components_to_use"] = [
            c for c in result.get("components_to_use", [])
            if c.lower() in valid_types
        ]
        if not result["components_to_use"]:
            result["components_to_use"] = ["info_card", "data_list"]
        elif result["components_to_use"] == ["markdown"]:
            result["components_to_use"] = ["info_card", "data_list", "markdown"]
        return result
    except json.JSONDecodeError:
        return {
            "components_to_use": ["info_card", "data_list", "markdown"],
            "component_descriptions": {
                "info_card": "摘要重點",
                "data_list": "結構化資訊",
                "markdown": "補充說明",
            },
        }


def call_openai_stage3_construct(text_content: dict, component_plan: dict) -> dict:
    """Stage 3: 構建最終的 AGUIResponse"""
    client = get_client()
    model  = st.secrets.get("OPENAI_MODEL", DEFAULT_MODEL)

    answer = text_content.get("answer", "")
    suggestions = text_content.get("suggestions", [])
    components_to_use = component_plan.get("components_to_use", ["markdown"])
    component_descriptions = component_plan.get("component_descriptions", {})

    system_prompt = textwrap.dedent("""
    你是一個 AG-UI 元件構建專家。根據計劃的 component 類型，將文本轉換為結構化 JSON。

    ## 可用 component 類型定義

    ### markdown
    {"type":"markdown","content":"Markdown 格式文字"}
    用於：一般說明、解釋、段落文字。

    ### info_card
    {"type":"info_card","title":"標題","description":"說明文字","variant":"info"}
    variant：info | warning | success | danger

    ### data_list
    {"type":"data_list","title":"標題（選填）","items":[{"label":"欄位","value":"內容"}]}

    ### step_process
    {"type":"step_process","title":"流程標題","steps":[{"title":"步驟1","description":"說明"}]}

    ### table
    {"type":"table","title":"表格標題","headers":["欄1","欄2"],"rows":[["A","B"]]}

    ### stat_grid
    {"type":"stat_grid","title":"統計標題","items":[{"label":"指標","value":"數值","description":"說明（選填）"}]}

    ### code_block
    {"type":"code_block","title":"標題（選填）","language":"python","content":"程式碼"}

    ### action_group
    {"type":"action_group","title":"動作標題","items":[{"label":"動作","action":"action_id","description":"說明"}]}

    ## 回傳格式
    必須回傳 JSON 物件，結構如下：
    {
      "components": [
        ... 一或多個 component 物件 ...
      ],
      "suggestions": ["建議1", "建議2"]
    }

    【重要】
    1. 回應必須是合法的 JSON，從 { 開始到 } 結束，無其他文字。
    2. suggestions 來自原始回答，直接傳入。
    3. 根據計劃的 component 類型構建 components 陣列。
    4. 優先使用結構化、可視化元件，不要把主要內容都塞進 markdown。
    5. 若答案含有步驟，轉成 step_process。
    6. 若答案含有重點整理、屬性、條件、注意事項，轉成 data_list 或 info_card。
    7. 若答案含有比較、分類、欄位整理，轉成 table。
    8. 若答案含有數值、統計、狀態摘要，轉成 stat_grid。
    9. markdown 最多只作為少量補充段落，避免單一 markdown component 承載全部內容。
    10. 繁體中文內容。
    11. 【禁止重複】每段內容只能出現在一個 component 中。若已放入 info_card.description，就不可再用 markdown 重複相同文字；若已放入 data_list 或 step_process，就不可再用 markdown 重複列出相同條目。每個 component 必須承載不同的資訊片段。
    """).strip()

    components_str = ", ".join(components_to_use)
    descriptions_str = "\n".join([f"- {k}: {v}" for k, v in component_descriptions.items()])

    user_message = f"""請根據以下計劃構建 AG-UI 回應：

計劃的 Component 類型：{components_str}

Component 使用指導：
{descriptions_str}

要轉換的答案文本：
{answer}

建議列表（保持不變）：
{json.dumps(suggestions, ensure_ascii=False)}
"""

    history = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message},
    ]

    resp = client.chat.completions.create(model=model, messages=history, temperature=0.5, max_tokens=4096)
    raw  = resp.choices[0].message.content

    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```\s*$", "", raw, flags=re.MULTILINE)
    raw = raw.strip()

    start = raw.find("{")
    end   = raw.rfind("}")
    if start != -1 and end != -1 and end > start:
        raw = raw[start:end + 1]

    try:
        result = json.loads(raw)
        # 確保 suggestions 與原始建議一致
        result["suggestions"] = suggestions
        return result
    except json.JSONDecodeError as e:
        return {
            "components": [
                {"type": "markdown", "content": answer}
            ],
            "suggestions": suggestions,
        }

def _extract_code_blocks(text: str) -> tuple[str, list[dict]]:
    code_blocks = []

    def replacer(match):
        language = (match.group(1) or "text").strip() or "text"
        content = match.group(2).strip()
        code_blocks.append({
            "type": "code_block",
            "title": "程式碼範例",
            "language": language,
            "content": content,
        })
        return ""

    cleaned = re.sub(r"```(\w+)?\n(.*?)```", replacer, text, flags=re.DOTALL)
    return cleaned.strip(), code_blocks

def _parse_step_process(text: str) -> dict | None:
    steps = []
    for line in text.splitlines():
        stripped = line.strip()
        match = re.match(r"^(?:\d+[\.\)]\s+)(.+)$", stripped)
        if match:
            content = match.group(1).strip()
            if ":" in content:
                title, desc = content.split(":", 1)
            elif "：" in content:
                title, desc = content.split("：", 1)
            else:
                title, desc = content, ""
            steps.append({"title": title.strip(), "description": desc.strip()})
    if len(steps) >= 2:
        return {"type": "step_process", "title": "操作流程", "steps": steps}
    return None

def _parse_data_list(text: str) -> dict | None:
    items = []
    for line in text.splitlines():
        stripped = line.strip().lstrip("-•").strip()
        if not stripped:
            continue
        if ":" in stripped:
            label, value = stripped.split(":", 1)
        elif "：" in stripped:
            label, value = stripped.split("：", 1)
        else:
            continue
        if label.strip() and value.strip():
            items.append({"label": label.strip(), "value": value.strip()})
    if len(items) >= 2:
        return {"type": "data_list", "title": "重點整理", "items": items[:8]}
    return None

def _parse_stat_grid(text: str) -> dict | None:
    items = []
    for line in text.splitlines():
        stripped = line.strip().lstrip("-•").strip()
        if not stripped:
            continue
        match = re.match(r"^([^:：]+)[:：]\s*([<>~=]?\s*[\d\.]+%?|[\d\.]+(?:\s*[A-Za-z]+)?)\s*(.*)$", stripped)
        if match:
            items.append({
                "label": match.group(1).strip(),
                "value": match.group(2).strip(),
                "description": match.group(3).strip() or None,
            })
    if 1 <= len(items) <= 4:
        return {"type": "stat_grid", "title": "摘要指標", "items": items}
    return None

def _build_info_card(text: str) -> dict | None:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        return None
    first = paragraphs[0]
    variant = "info"
    lowered = first.lower()
    if any(word in first for word in ["注意", "警告", "風險", "錯誤"]):
        variant = "warning"
    elif any(word in first for word in ["可以", "建議", "適合", "成功"]):
        variant = "success"
    return {
        "type": "info_card",
        "title": "重點摘要",
        "description": first[:240],
        "variant": variant,
    }

def _build_markdown_remainder(text: str) -> dict | None:
    cleaned = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not cleaned:
        return None
    return {"type": "markdown", "content": cleaned}

def normalize_agui_response(agui_resp: dict, fallback_answer: str) -> dict:
    components = agui_resp.get("components", []) or []
    if not components:
        components = [{"type": "markdown", "content": fallback_answer}]

    non_markdown = [c for c in components if c.get("type") != "markdown"]
    markdown_parts = [
        c.get("content", "").strip()
        for c in components
        if c.get("type") == "markdown" and c.get("content", "").strip()
    ]

    if non_markdown and markdown_parts:
        # 已有結構化元件時，只保留真正補充性的 markdown（短段落）
        # 避免 Stage 3 把相同內容同時塞進 info_card.description 和 markdown.content
        filtered_markdown = [
            c for c in components
            if c.get("type") != "markdown" or (
                len(c.get("content", "").strip()) <= 280
                and not any(
                    c.get("content", "").strip()[:80] in other.get("description", "")
                    or c.get("content", "").strip()[:80] in other.get("content", "")
                    for other in non_markdown
                )
            )
        ]
        agui_resp["components"] = filtered_markdown
        return agui_resp

    if non_markdown:
        agui_resp["components"] = components
        return agui_resp

    source_text = "\n\n".join(markdown_parts) or fallback_answer
    text_without_code, code_blocks = _extract_code_blocks(source_text)

    rebuilt = []
    info_card = _build_info_card(text_without_code)
    if info_card:
        rebuilt.append(info_card)

    step_process = _parse_step_process(text_without_code)
    if step_process:
        rebuilt.append(step_process)

    stat_grid = _parse_stat_grid(text_without_code)
    if stat_grid:
        rebuilt.append(stat_grid)

    data_list = _parse_data_list(text_without_code)
    if data_list:
        rebuilt.append(data_list)

    rebuilt.extend(code_blocks)

    markdown_remainder = _build_markdown_remainder(text_without_code)
    if markdown_remainder and len(rebuilt) < 2:
        rebuilt.append(markdown_remainder)
    elif markdown_remainder and not any(c.get("type") == "markdown" for c in rebuilt):
        if len(markdown_remainder["content"]) <= 280:
            rebuilt.append(markdown_remainder)

    agui_resp["components"] = rebuilt or [{"type": "markdown", "content": fallback_answer}]
    return agui_resp

def extract_markdown_artifacts(markdown_text: str) -> list[dict]:
    artifacts = []
    for idx, match in enumerate(re.finditer(r"```(\w+)?\n(.*?)```", markdown_text, flags=re.DOTALL)):
        language = (match.group(1) or "").strip()
        content = match.group(2).strip()
        detected = detect_artifact_kind(content, language)
        if detected:
            ext, normalized = detected
            artifacts.append({
                "ext": ext,
                "content": normalized,
                "title": f"artifact_{idx + 1}",
            })

    if not artifacts:
        direct = detect_artifact_kind(markdown_text)
        if direct:
            ext, normalized = direct
            artifacts.append({"ext": ext, "content": normalized, "title": "artifact"})

    return artifacts

def extract_artifacts_from_components(components: list) -> list[dict]:
    artifacts = []
    for idx, comp in enumerate(components):
        comp_type = comp.get("type", "")
        title = comp.get("title", f"artifact_{idx + 1}")

        if comp_type == "code_block":
            detected = detect_artifact_kind(comp.get("content", ""), comp.get("language", ""))
            if detected:
                ext, normalized = detected
                artifacts.append({"ext": ext, "content": normalized, "title": title})
        elif comp_type == "markdown":
            artifacts.extend(extract_markdown_artifacts(comp.get("content", "")))

    return artifacts

def persist_component_artifacts(components: list):
    artifacts = extract_artifacts_from_components(components)
    created = []
    for artifact in artifacts:
        name = make_artifact_name(artifact["ext"], artifact.get("title", "artifact"))
        meta = upsert_generated_artifact(
            name=name,
            content=artifact["content"],
            ext=artifact["ext"],
            prompt=artifact.get("title", ""),
        )
        created.append(meta)
    return created


# ══════════════════════════════════════════════════════════════
# 呼叫 OpenAI → AGUIResponse（三階段編排）
# ══════════════════════════════════════════════════════════════
def call_openai(user_text: str, attached_file: dict | None = None, progress_target=None) -> dict:
    """三階段生成 AG-UI 回應：(1) 生成文本 (2) 決定 components (3) 構建最終回應"""
    if get_client() is None:
        return {
            "components": [
                {"type": "info_card", "title": "設定錯誤", "description": "請先設定您的 OpenAI API Key 以開始對話。", "variant": "danger"}
            ],
            "suggestions": [],
        }
    try:
        st.session_state.thinking_stages = []

        # Stage 1: 生成純文本內容和建議
        st.session_state.thinking_stages.append("📝 Stage 1: 生成內容...")
        if progress_target is not None:
            render_thinking_bubble(progress_target, st.session_state.thinking_stages)
        text_content = call_openai_stage1_text_content(user_text, attached_file)

        # Stage 2: 決定 component 計劃
        st.session_state.thinking_stages.append("🎨 Stage 2: 決定元件類型...")
        if progress_target is not None:
            render_thinking_bubble(progress_target, st.session_state.thinking_stages)
        component_plan = call_openai_stage2_component_plan(text_content.get("answer", ""))

        # Stage 3: 構建最終回應
        st.session_state.thinking_stages.append("🔨 Stage 3: 構建最終回應...")
        if progress_target is not None:
            render_thinking_bubble(progress_target, st.session_state.thinking_stages)
        agui_resp = call_openai_stage3_construct(text_content, component_plan)
        agui_resp = normalize_agui_response(agui_resp, text_content.get("answer", ""))
        persist_component_artifacts(agui_resp.get("components", []))

        # 檢查是否有 surface component
        has_surface = any(c.get("type") == "surface" for c in agui_resp.get("components", []))
        if has_surface:
            if not st.session_state.stage_expanded:
                st.session_state.stage_expanded = True
                st.session_state.split_ratio = 25

        return agui_resp

    except Exception as e:
        st.error(f"生成回應時出錯：{e}")
        return {
            "components": [
                {"type": "info_card", "title": "處理錯誤", "description": f"無法生成回應：{str(e)[:200]}", "variant": "danger"}
            ],
            "suggestions": [],
        }

def pdf_extract_text(path: Path) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            return "\n\n".join(page.extract_text() or "" for page in pdf.pages[:10])
    except Exception as e:
        return f"（PDF 解析失敗：{e}）"

def docx_extract_text(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"（DOCX 解析失敗：{e}）"

# ══════════════════════════════════════════════════════════════
# 渲染 AG-UI 元件
# ══════════════════════════════════════════════════════════════
def render_agui_components(components: list):
    for comp in components:
        t = comp.get("type", "")

        # ── markdown ──
        if t == "markdown":
            st.markdown(comp.get("content", ""))

        # ── info_card ──
        elif t == "info_card":
            variant = comp.get("variant", "info")
            # 大地色系配色
            color_map = {
                "info": {"border": "#0f766e", "bg": "#d7efe8"},
                "warning": {"border": "#b45309", "bg": "#fff8eb"},
                "success": {"border": "#15803d", "bg": "#dcfce7"},
                "danger": {"border": "#dc2626", "bg": "#fee2e2"}
            }
            colors = color_map.get(variant, color_map["info"])
            st.markdown(f"""
<div style="border-left:4px solid {colors['border']}; padding:14px; background:{colors['bg']}; border: 1px solid rgba(216, 205, 189, 0.4); margin:12px 0; border-radius:16px;">
  <div style="font-weight:700; font-size:1em; color:{colors['border']}; margin-bottom:6px;">{comp.get('title','')}</div>
  <div style="color:#1d2a2a; font-size:13px; line-height:1.6;">{comp.get('description','')}</div>
</div>""", unsafe_allow_html=True)

        # ── data_list ──
        elif t == "data_list":
            if comp.get("title"):
                st.caption(comp["title"])
            with st.container(border=True):
                for i, item in enumerate(comp.get("items", [])):
                    if i > 0:
                        st.markdown("<hr style='margin: 8px 0; border: none; border-top: 1px solid rgba(128, 128, 128, 0.15);'>", unsafe_allow_html=True)
                    lbl = item.get("label","")
                    val = item.get("value","")
                    left, right = st.columns([1, 3], gap="small")
                    with left:
                        st.markdown(f"**{lbl}**")
                    with right:
                        st.write(val)

        # ── step_process ──
        elif t == "step_process":
            st.markdown(f"### {comp.get('title','流程')}")
            with st.container(border=True):
                for i, step in enumerate(comp.get("steps",[]), 1):
                    if i > 1:
                        st.markdown("<hr style='margin: 12px 0; border: none; border-top: 1px solid rgba(128, 128, 128, 0.15);'>", unsafe_allow_html=True)
                    title = step.get("title","")
                    desc  = step.get("description","")
                    ncol, ccol = st.columns([1, 12], gap="small")
                    with ncol:
                        st.markdown(f"<div style='text-align:center; font-weight:bold; font-size:1.1em;'>{i}</div>", unsafe_allow_html=True)
                    with ccol:
                        st.markdown(f"**{title}**")
                        if desc:
                            st.write(desc)

        # ── table ──
        elif t == "table":
            if comp.get("title"):
                st.markdown(f"**{comp['title']}**")
            headers = comp.get("headers", [])
            rows    = comp.get("rows", [])
            if headers and rows:
                import pandas as pd
                df = pd.DataFrame(rows, columns=headers)
                st.dataframe(df, use_container_width=True)

        # ── stat_grid ──
        elif t == "stat_grid":
            if comp.get("title"):
                st.markdown(f"**{comp['title']}**")
            items = comp.get("items", [])
            cols  = st.columns(len(items))
            for ci, item in enumerate(items):
                with cols[ci]:
                    st.metric(label=item.get("label",""), value=item.get("value",""),
                              help=item.get("description"))

        # ── code_block ──
        elif t == "code_block":
            if comp.get("title"):
                st.markdown(f"**{comp['title']}**")
            lang = comp.get("language", "python")
            code = comp.get("content", "")
            st.code(code, language=lang)

        # ── action_group ──
        elif t == "action_group":
            st.markdown(f"**{comp.get('title','建議動作')}**")
            items = comp.get("items", [])
            cols = st.columns(len(items)) if items else []
            for idx, item in enumerate(items):
                lbl = item.get("label", "")
                act = item.get("action", "")
                desc = item.get("description", "")
                btn_text = f"{lbl}"
                if desc:
                    btn_text += f" - {desc}"
                with cols[idx]:
                    if st.button(btn_text, key=f"act_{act}_{id(comp)}", use_container_width=True):
                        st.session_state.pending_suggestion = lbl
                        st.rerun()

        # ── surface ──
        elif t == "surface":
            kind  = comp.get("kind", "html")
            title = comp.get("title", "surface")
            if kind == "html":
                html_src = comp.get("html", "")
                css      = comp.get("css", "")
                full     = f"<style>{css}</style>{html_src}"
                fname    = f"{title.replace(' ','_')}.html"
                # 用 prompt（title）去重，避免每次 rerun 重複寫入
                existing = next(
                    (f for f in st.session_state.generated_files
                     if f.get("prompt") == title and f["ext"] == "html"),
                    None,
                )
                if existing is None:
                    meta = save_generated(fname, full.encode("utf-8"), "html", prompt=title)
                    st.session_state.generated_files.append(meta)
                    st.session_state.active_file = meta["id"]
                st.success(f"✅ 已產生 surface：{title}，請至右側展示介面查看。")
            elif kind == "svg":
                svg_src = comp.get("svg", "")
                fname   = f"{title.replace(' ','_')}.svg"
                existing = next(
                    (f for f in st.session_state.generated_files
                     if f.get("prompt") == title and f["ext"] == "svg"),
                    None,
                )
                if existing is None:
                    meta = save_generated(fname, svg_src.encode("utf-8"), "svg", prompt=title)
                    st.session_state.generated_files.append(meta)
                    st.session_state.active_file = meta["id"]
                st.success(f"✅ 已產生 SVG：{title}，請至右側展示介面查看。")
            else:
                st.info(f"surface kind={kind} 暫不支援")

        # ── 未知類型 ──
        else:
            st.warning(f"未知 component type: {t}")

def render_thinking_bubble(target, stage_lines: list[str]):
    stage_html = "".join(
        f'<div style="margin-top:8px;">{line}</div>' for line in stage_lines
    )
    target.markdown(f"""
<div class="msg-meta">Agent</div>
<div class="thinking-bubble">
  <div style="display:flex;align-items:center;">
    <div class="thinking-dot"></div>
    <div class="thinking-dot"></div>
    <div class="thinking-dot"></div>
    <span style="margin-left:4px;">思考中…</span>
  </div>
  <div style="margin-top:6px; width:100%;">{stage_html}</div>
</div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
# 自定義 CSS（溫暖大地色系 - 奶茶色調）
# ══════════════════════════════════════════════════════════════
st.markdown("""
<style>
:root {
    --bg: #f4efe6;
    --panel: rgba(255, 250, 242, 0.92);
    --panel-strong: #fffdf8;
    --ink: #1d2a2a;
    --muted: #61706f;
    --line: #d8cdbd;
    --accent: #0f766e;
    --accent-soft: #d7efe8;
    --warn: #b45309;
    --shadow: 0 20px 50px rgba(60, 42, 20, 0.12);
    --user-msg: #1f2937;
}

* {
    margin: 0 !important;
    padding: 0 !important;
    box-sizing: border-box !important;
}

/* 隱藏 Streamlit header */
header[data-testid="stHeader"] {
    display: none !important;
}

/* 全局背景 - 雙層徑向漸層 + 基底色 */
html, body {
    background: radial-gradient(circle at top left, rgba(15, 118, 110, 0.12), transparent 28%),
                radial-gradient(circle at bottom right, rgba(180, 83, 9, 0.08), transparent 24%),
                var(--bg) !important;
    font-family: "Segoe UI", "Noto Sans TC", sans-serif !important;
    color: var(--ink) !important;
}

html, body, [data-testid="stAppViewContainer"], .main, .main .block-container, div[data-testid="stHorizontalBlock"], div[data-testid="column"], div[data-testid="stVerticalBlock"] {
    margin: 0 !important;
    padding: 0 !important;
    width: 100% !important;
    min-width: 100% !important;
    height: 100% !important;
    min-height: 100% !important;
    overflow: hidden !important;
}

[data-testid="stAppViewContainer"] {
    display: flex !important;
    flex-direction: column !important;
    gap: 0 !important;
}

.main, .main .block-container, div[data-testid="stHorizontalBlock"], div[data-testid="column"], div[data-testid="stVerticalBlock"] {
    display: flex !important;
    flex-direction: column !important;
    flex: 1 !important;
    gap: 0 !important;
}

div[data-testid="stHorizontalBlock"] {
    flex-direction: row !important;
}

section[data-testid="stSidebar"] {
    display: none !important;
}

/* 訊息氣泡文字換行，避免橫向溢出 */
.msg-user, .msg-agent {
    word-break: break-word !important;
    overflow-wrap: break-word !important;
    max-width: 100% !important;
    box-sizing: border-box !important;
}

/* 所有 agent bubble 內容也限制寬度 */
.agent-bubble {
    max-width: 100% !important;
    overflow-x: hidden !important;
    word-break: break-word !important;
}

/* stForm 不要超出寬度 */
div[data-testid="stForm"] {
    overflow-x: hidden !important;
}

/* 對話串獨立垂直滾動容器 — 高度由 JS 動態計算 */
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .chat-scroll-marker) {
    min-height: 0 !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    padding-right: 6px;
}

/* 側邊欄檔案列表獨立垂直滾動容器 */
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .sb-scroll-marker) {
    max-height: calc(100vh - 250px) !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    padding-right: 6px;
}

/* 展示區 Tab 內容獨立垂直滾動容器 */
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .tab-scroll-marker) {
    max-height: calc(100vh - 230px) !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    padding-right: 6px;
}

/* 自定義滾動條 - 大地色系 */
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .chat-scroll-marker)::-webkit-scrollbar,
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .sb-scroll-marker)::-webkit-scrollbar,
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .tab-scroll-marker)::-webkit-scrollbar {
    width: 6px;
}

div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .chat-scroll-marker)::-webkit-scrollbar-track,
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .sb-scroll-marker)::-webkit-scrollbar-track,
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .tab-scroll-marker)::-webkit-scrollbar-track {
    background: rgba(216, 205, 189, 0.2);
    border-radius: 3px;
}

div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .chat-scroll-marker)::-webkit-scrollbar-thumb,
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .sb-scroll-marker)::-webkit-scrollbar-thumb,
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .tab-scroll-marker)::-webkit-scrollbar-thumb {
    background: rgba(15, 118, 110, 0.3);
    border-radius: 3px;
}

div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .chat-scroll-marker)::-webkit-scrollbar-thumb:hover,
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .sb-scroll-marker)::-webkit-scrollbar-thumb:hover,
div[data-testid="stVerticalBlock"]:has(> div[data-testid="element-container"] .tab-scroll-marker)::-webkit-scrollbar-thumb:hover {
    background: rgba(15, 118, 110, 0.5);
}

/* Panel Header - 大地色系漸層 */
.panel-hdr {
    background: linear-gradient(135deg, #0f766e, #115e59);
    color: white;
    padding: 18px 24px;
    border-bottom: 1px solid var(--line);
    font-weight: 600;
    margin-bottom: 0;
    display: flex;
    justify-content: space-between;
    align-items: center;
}

.panel-hdr h3 {
    margin: 0;
    font-size: 1.1em;
}

/* 訊息標籤 */
.msg-meta {
    font-size: 11px;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: var(--muted);
    margin-bottom: 8px;
    font-weight: 700;
}

/* User 訊息氣泡 - 深灰色 */
.msg-user {
    align-self: flex-end;
    background: var(--user-msg);
    color: white;
    padding: 14px 16px;
    border-radius: 20px;
    border-bottom-right-radius: 8px;
    margin-bottom: 12px;
    max-width: 92%;
    box-shadow: 0 8px 24px rgba(52, 39, 23, 0.06);
}

/* Agent 訊息氣泡 - 奶白色 */
.msg-agent {
    background: var(--panel-strong);
    color: var(--ink);
    padding: 14px 18px;
    border-radius: 20px;
    border-bottom-left-radius: 8px;
    margin-bottom: 12px;
    border: 1px solid rgba(216, 205, 189, 0.85);
    box-shadow: 0 8px 24px rgba(52, 39, 23, 0.06);
}

.agent-bubble {
    margin-bottom: 12px;
}

.agent-bubble div[data-testid="stVerticalBlockBorderWrapper"] {
    background: var(--panel-strong);
    border: 1px solid rgba(216, 205, 189, 0.85);
    border-radius: 20px;
    border-bottom-left-radius: 8px;
    padding: 14px 18px;
}

/* 檔案標籤 */
.file-chip {
    display: inline-flex;
    align-items: center;
    background: rgba(216, 205, 189, 0.25);
    color: var(--ink);
    padding: 4px 10px;
    border-radius: 12px;
    font-size: 0.9em;
    margin-top: 6px;
    border: 1px solid rgba(216, 205, 189, 0.4);
}

/* 思考氣泡 */
.thinking-bubble {
    background: var(--panel-strong);
    color: var(--ink);
    padding: 12px 18px;
    border-radius: 20px;
    display: inline-flex;
    flex-direction: column;
    align-items: flex-start;
    border: 1px solid rgba(216, 205, 189, 0.85);
    margin-bottom: 12px;
    box-shadow: 0 8px 24px rgba(52, 39, 23, 0.06);
}

.thinking-dot {
    width: 8px;
    height: 8px;
    background: var(--accent);
    border-radius: 50%;
    margin: 0 3px;
    animation: thinking 1.4s infinite ease-in-out both;
}

.thinking-dot:nth-child(1) { animation-delay: -0.32s; }
.thinking-dot:nth-child(2) { animation-delay: -0.16s; }

@keyframes thinking {
  0%, 80%, 100% { transform: scale(0.8); opacity: 0.5; }
  40% { transform: scale(1); opacity: 1; }
}

/* PDF 預覽框 */
.pdf-frame {
    border: 1px solid var(--line);
    border-radius: 14px;
    background: white;
}

/* 折疊按鈕 */
.toggle-btn {
    background: rgba(15, 118, 110, 0.1);
    border: 1px solid rgba(15, 118, 110, 0.2);
    color: var(--accent);
    padding: 4px 12px;
    border-radius: 999px;
    cursor: pointer;
    font-size: 0.9em;
    transition: all 0.2s;
}

.toggle-btn:hover {
    background: rgba(15, 118, 110, 0.2);
}

/* Streamlit 元件覆寫 */
div[data-testid="stVerticalBlockBorderWrapper"] {
    background: var(--panel) !important;
    border: 1px solid var(--line) !important;
    border-radius: 16px !important;
}

/* 按鈕樣式覆寫 */
.stButton button {
    background: var(--accent) !important;
    color: white !important;
    border: none !important;
    border-radius: 14px !important;
    font-weight: 600 !important;
    transition: all 0.2s !important;
}

.stButton button:hover {
    background: #115e59 !important;
    box-shadow: 0 4px 12px rgba(15, 118, 110, 0.25) !important;
}

.stButton button[kind="secondary"] {
    background: rgba(216, 205, 189, 0.3) !important;
    color: var(--ink) !important;
}

.stButton button[kind="secondary"]:hover {
    background: rgba(216, 205, 189, 0.5) !important;
}

/* 輸入框樣式 */
input, textarea {
    background: white !important;
    border: 1px solid var(--line) !important;
    border-radius: 14px !important;
    color: var(--ink) !important;
}

input:focus, textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 2px rgba(15, 118, 110, 0.1) !important;
}

/* Tab 樣式 */
.stTabs [data-baseweb="tab-list"] {
    gap: 10px;
}

.stTabs [data-baseweb="tab"] {
    background: rgba(255, 255, 255, 0.7) !important;
    border: 1px solid var(--line) !important;
    border-radius: 14px !important;
    color: var(--muted) !important;
    padding: 10px 14px !important;
}

.stTabs [aria-selected="true"] {
    background: var(--accent-soft) !important;
    border-color: rgba(15, 118, 110, 0.25) !important;
    color: #0b5c55 !important;
    font-weight: 700 !important;
}

/* Metric 卡片 */
div[data-testid="stMetric"] {
    background: white !important;
    border: 1px solid var(--line) !important;
    border-radius: 18px !important;
    padding: 16px !important;
}

/* Dataframe 樣式 */
div[data-testid="stDataFrame"] {
    border: 1px solid var(--line) !important;
    border-radius: 14px !important;
    overflow: hidden !important;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<script>
(function() {
  function fitChatMessageScroll() {
    var marker = document.querySelector('.chat-scroll-marker');
    var formMarker = document.querySelector('.chat-form-marker');
    if (!marker || !formMarker) return;

    var scrollBlock = marker.closest('[data-testid="stVerticalBlock"]');
    var columnEl = formMarker.closest('[data-testid="column"]');
    if (!scrollBlock || !columnEl) return;

    var formEl = columnEl.querySelector('[data-testid="stForm"]');
    var targetTop = null;

    if (formEl) {
      targetTop = formEl.getBoundingClientRect().top;
    } else {
      var formBlock = formMarker.closest('[data-testid="stVerticalBlock"]');
      if (!formBlock) return;
      targetTop = formBlock.getBoundingClientRect().top;
    }

    var scrollTop = scrollBlock.getBoundingClientRect().top;
    var available = targetTop - scrollTop - 12;
    if (available < 180) available = 180;

    scrollBlock.style.setProperty('height', available + 'px', 'important');
    scrollBlock.style.setProperty('max-height', available + 'px', 'important');
  }

  var observer = new MutationObserver(function() {
    requestAnimationFrame(fitChatMessageScroll);
  });

  function init() {
    fitChatMessageScroll();
    observer.observe(document.body, { childList: true, subtree: true });
    window.addEventListener('resize', fitChatMessageScroll);
  }

  if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', init);
  } else {
    setTimeout(init, 150);
  }
})();
</script>
""", unsafe_allow_html=True)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 主版面配置
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# 根據展開狀態計算列寬比例
if not st.session_state.chat_expanded and not st.session_state.stage_expanded:
    # 兩邊都折疊：顯示小寬度（各占 50%）
    col_widths = [1, 1]
elif not st.session_state.chat_expanded:
    # 對話欄折疊，展示欄展開
    col_widths = [1, 9]
elif not st.session_state.stage_expanded:
    # 對話欄展開，展示欄折疊
    col_widths = [9, 1]
else:
    # 兩邊都展開：使用自定義比例
    chat_ratio = st.session_state.split_ratio
    stage_ratio = 100 - chat_ratio
    col_widths = [chat_ratio, stage_ratio]

chat_col, stage_col = st.columns(col_widths, gap="small")
thinking_placeholder = None

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 左欄：對話流
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
with chat_col:
    # Panel Header with Toggle Button
    toggle_label = "◀" if st.session_state.chat_expanded else "▶"
    
    col_title, col_btn = st.columns([4, 1])
    with col_title:
        st.markdown('<div class="panel-hdr"><h3>對話流</h3></div>', unsafe_allow_html=True)
    with col_btn:
        if st.button(toggle_label, key="chat_toggle", help="展開/折疊對話欄"):
            st.session_state.chat_expanded = not st.session_state.chat_expanded
            st.rerun()

    if st.session_state.chat_expanded:
        with st.container():
            st.markdown('<div class="chat-body-marker"></div>', unsafe_allow_html=True)
            st.markdown('<div class="chat-container">', unsafe_allow_html=True)

            chat_scroll_container = st.container()
            with chat_scroll_container:
                st.markdown('<div class="chat-scroll-marker"></div>', unsafe_allow_html=True)
                for i, m in enumerate(st.session_state.messages):
                    if m["role"] == "user":
                        st.markdown('<div class="msg-meta">User</div>', unsafe_allow_html=True)
                        content_text = m.get("content", "")
                        chip = m.get("file_chip")
                        chip_html = ""
                        if chip:
                            chip_html = f'<div class="file-chip">{file_icon(chip["ext"])} {chip["name"]}</div>'
                        st.markdown(
                            f'<div class="msg-user">{content_text}{chip_html}</div>',
                            unsafe_allow_html=True,
                        )

                    else:
                        st.markdown('<div class="msg-meta">Agent</div>', unsafe_allow_html=True)
                        st.markdown('<div class="agent-bubble">', unsafe_allow_html=True)
                        with st.container(border=True):
                            agui = m.get("agui")
                            if agui:
                                render_agui_components(agui.get("components", []))
                                suggestions = agui.get("suggestions", [])
                                if suggestions:
                                    st.markdown('<div style="display:flex;flex-wrap:wrap;gap:6px;margin-top:10px;">', unsafe_allow_html=True)
                                    sugg_cols = st.columns(len(suggestions))
                                    for si, sugg_text in enumerate(suggestions):
                                        with sugg_cols[si]:
                                            if st.button(sugg_text, key=f"sugg_{id(agui)}_{si}",
                                                         use_container_width=True):
                                                st.session_state.pending_suggestion = sugg_text
                                                st.rerun()
                                    st.markdown('</div>', unsafe_allow_html=True)
                            else:
                                st.markdown(m.get("content",""))
                        st.markdown('</div>', unsafe_allow_html=True)

                thinking_placeholder = st.empty()
                stop_button_placeholder = st.empty()
                if st.session_state.thinking:
                    render_thinking_bubble(thinking_placeholder, st.session_state.get("thinking_stages", []))
                    with stop_button_placeholder.container():
                        if st.button("🛑 停止生成", key="stop_btn", use_container_width=True):
                            st.session_state.thinking = False
                            st.session_state.thinking_stages = []
                            st.rerun()

                guide_container = st.empty()
                if not st.session_state.messages:
                    with guide_container.container():
                        st.markdown("""
                        <div style="text-align: center; padding: 20px 10px; margin-bottom: 10px;">
                          <h2 style="color: #0f766e; margin-bottom: 8px; font-weight: 600;">歡迎使用 Agent UI/UX 平台 🤖</h2>
                          <p style="color: #61706f; font-size: 0.95em;">我是一個支援多種視覺化元件與 Artifact 生成的 AI 助手。點選下方範例或直接輸入指令開始對話：</p>
                        </div>
                        """, unsafe_allow_html=True)

                        examples = [
                            "📊 比較 React, Vue 和 Angular 的優缺點",
                            "🧭 設計一個三步驟的軟體部署流程說明",
                            "📈 產生伺服器 CPU 和記憶體指標的數據看板",
                            "🎨 畫一個漂亮的 SVG 圓餅圖"
                        ]

                        cols1 = st.columns(2)
                        cols2 = st.columns(2)
                        for idx, ex in enumerate(examples):
                            target_col = cols1[idx] if idx < 2 else cols2[idx - 2]
                            with target_col:
                                if st.button(ex, key=f"ex_btn_{idx}", use_container_width=True):
                                    st.session_state.pending_suggestion = ex
                                    st.rerun()

            st.markdown('</div>', unsafe_allow_html=True)

        has_api_key = bool(st.secrets.get("OPENAI_API_KEY", "") or st.session_state.get("api_key_override", ""))
        with st.container():
            st.markdown('<div class="chat-form-marker"></div>', unsafe_allow_html=True)
            if not has_api_key:
                with st.container(border=True):
                    st.markdown("""
                    <div style="padding: 5px; border-radius: 8px;">
                      <h4 style="margin-top:0; color:#0f766e; font-weight: 600;">🔑 設定 OpenAI API 金鑰</h4>
                      <p style="font-size:0.9em; color:#61706f; margin-bottom: 12px;">本機環境未設定 API Key。請在下方輸入您的 OpenAI API Key，金鑰僅會暫存在您的瀏覽器會話中。</p>
                    </div>
                    """, unsafe_allow_html=True)
                    user_key = st.text_input("輸入 API 金鑰 (API Key)", type="password", placeholder="sk-...", label_visibility="collapsed")
                    if st.button("確認儲存金鑰", use_container_width=True):
                        if user_key.strip():
                            st.session_state.api_key_override = user_key.strip()
                            st.session_state.client = None
                            st.success("金鑰設定成功！")
                            st.rerun()
                        else:
                            st.error("請輸入有效的金鑰！")
            else:
                prefill = st.session_state.get("pending_suggestion", "")
                with st.form("composer", clear_on_submit=True):
                    attached = st.file_uploader(
                        "附加檔案",
                        type=["pdf","docx","doc","png","jpg","jpeg","gif","webp","txt","md"],
                        label_visibility="collapsed",
                        disabled=st.session_state.thinking,
                    )
                    user_input = st.text_input(
                        "輸入指令", value=prefill,
                        placeholder="輸入指令，或上傳檔案後送出…",
                        label_visibility="collapsed",
                        disabled=st.session_state.thinking,
                    )
                    submitted = st.form_submit_button("⬆ Send", use_container_width=True, disabled=st.session_state.thinking)

                if submitted and (user_input.strip() or attached):
                    final_text = user_input.strip()
                    st.session_state.pending_suggestion = ""

                    attached_meta = None
                    if attached:
                        attached_meta = save_upload(attached)
                        if not any(f["id"] == attached_meta["id"] for f in st.session_state.uploaded_files):
                            st.session_state.uploaded_files.append(attached_meta)
                        st.session_state.active_file = attached_meta["id"]

                    chip = {"name": attached_meta["name"], "ext": attached_meta["ext"]} if attached_meta else None

                    st.session_state.messages.append({
                        "role": "user",
                        "content": final_text or f"（已上傳 {attached.name}）",
                        "file_chip": chip,
                    })

                    st.session_state.thinking_stages = []
                    st.session_state.thinking = True
                    st.rerun()

# ── 在 thinking 狀態下執行 API 呼叫 ──
if st.session_state.thinking:
    last_user = next(
        (m for m in reversed(st.session_state.messages) if m["role"] == "user"), None
    )
    attached_meta = find_file(st.session_state.active_file) if st.session_state.active_file else None
    # 只在剛上傳的那次傳入 attached（避免每次 rerun 都重傳）
    is_upload_msg = (
        attached_meta and
        attached_meta.get("source") == "upload" and
        last_user and
        last_user.get("file_chip") and
        last_user["file_chip"]["name"] == attached_meta["name"]
    )

    agui_resp = call_openai(
        last_user["content"] if last_user else "",
        attached_meta if is_upload_msg else None,
        progress_target=thinking_placeholder,
    )

    st.session_state.messages.append({
        "role": "assistant",
        "agui": agui_resp,
    })
    st.session_state.thinking = False
    st.session_state.thinking_stages = []
    st.rerun()

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 右欄：展示介面
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
with stage_col:
    # Panel Header with Toggle Button and Resize Slider
    toggle_label = "▶" if st.session_state.stage_expanded else "◀"
    
    col_title2, col_btn2 = st.columns([4, 1])
    with col_title2:
        st.markdown('<div class="panel-hdr"><h3>展示介面</h3></div>', unsafe_allow_html=True)
    with col_btn2:
        if st.button(toggle_label, key="stage_toggle", help="展開/折疊展示欄"):
            st.session_state.stage_expanded = not st.session_state.stage_expanded
            st.rerun()

    if st.session_state.stage_expanded:
        # 比例調整滑桿（只有當兩邊都展開時才顯示）
        if st.session_state.chat_expanded:
            cols_ratio = st.columns([1, 4])
            with cols_ratio[0]:
                st.markdown("<span style='font-size:0.9em;color:#888;line-height:2.8;'>🌓 調整畫面比例:</span>", unsafe_allow_html=True)
            with cols_ratio[1]:
                new_ratio = st.slider(
                    "對話欄占比 (%)",
                    min_value=10,
                    max_value=90,
                    value=st.session_state.split_ratio,
                    step=5,
                    key="ratio_slider",
                    label_visibility="collapsed"
                )
                if new_ratio != st.session_state.split_ratio:
                    st.session_state.split_ratio = new_ratio
                    st.rerun()

        sb_col, main_col = st.columns([2, 5], gap="small")

        # ── Sidebar ──
        with sb_col:
            st.markdown("**🔍 搜尋與篩選**")
            search_q = st.text_input("搜尋檔名...", key="file_search", label_visibility="collapsed")
            
            sb_scroll_container = st.container()
            with sb_scroll_container:
                st.markdown('<div class="sb-scroll-marker"></div>', unsafe_allow_html=True)
                st.markdown("**上傳的檔案**")
                filtered_uploads = [
                    f for f in st.session_state.uploaded_files
                    if not search_q or search_q.lower() in f["name"].lower()
                ]
                if not filtered_uploads:
                    st.caption("無符合檔案")
                else:
                    for f in filtered_uploads:
                        is_active = st.session_state.active_file == f["id"]
                        file_cols = st.columns([5, 1])
                        with file_cols[0]:
                            if st.button(f"{file_icon(f['ext'])} {f['name']}", key=f"ub_{f['id']}",
                                         use_container_width=True,
                                         type="primary" if is_active else "secondary"):
                                st.session_state.active_file = f["id"]
                                st.rerun()
                        with file_cols[1]:
                            if st.button("🗑️", key=f"del_ub_{f['id']}", help="刪除此檔案", use_container_width=True):
                                try:
                                    if Path(f["path"]).exists():
                                        Path(f["path"]).unlink()
                                except:
                                    pass
                                st.session_state.uploaded_files.remove(f)
                                if st.session_state.active_file == f["id"]:
                                    st.session_state.active_file = None
                                st.rerun()

                new_file = st.file_uploader(
                    "新增上傳",
                    type=["pdf","docx","doc","png","jpg","jpeg","gif","webp","txt","md"],
                    key="sb_upload", label_visibility="collapsed",
                )
                if new_file:
                    meta = save_upload(new_file)
                    if not any(f["name"] == meta["name"] for f in st.session_state.uploaded_files):
                        st.session_state.uploaded_files.append(meta)
                        st.session_state.active_file = meta["id"]
                        st.rerun()

                st.divider()

                st.markdown("**模型產出**")
                filtered_generated = [
                    f for f in st.session_state.generated_files
                    if not search_q or search_q.lower() in f["name"].lower()
                ]
                if not filtered_generated:
                    st.caption("無符合檔案")
                else:
                    for f in filtered_generated:
                        is_active = st.session_state.active_file == f["id"]
                        file_cols = st.columns([5, 1])
                        with file_cols[0]:
                            if st.button(f"{file_icon(f['ext'])} {f['name']}", key=f"gb_{f['id']}",
                                         use_container_width=True,
                                         type="primary" if is_active else "secondary"):
                                st.session_state.active_file = f["id"]
                                st.rerun()
                        with file_cols[1]:
                            if st.button("🗑️", key=f"del_gb_{f['id']}", help="刪除此檔案", use_container_width=True):
                                try:
                                    if Path(f["path"]).exists():
                                        Path(f["path"]).unlink()
                                except:
                                    pass
                                st.session_state.generated_files.remove(f)
                                if st.session_state.active_file == f["id"]:
                                    st.session_state.active_file = None
                                st.rerun()

        # ── Main Viewer ──
        with main_col:
            active = find_file(st.session_state.active_file) if st.session_state.active_file else None

            if active is None:
                st.info("從左側選擇檔案，或透過對話要求 Agent 生成內容。")
            else:
                ext  = active["ext"].lower()
                path = Path(active["path"])
                src  = active["source"]

                badge = "🔵 上傳" if src == "upload" else "🟢 模型產出"
                st.markdown(
                    f"**{file_icon(ext)} {active['name']}** &nbsp; `{badge}` &nbsp; `{active['ts']}`",
                    unsafe_allow_html=True,
                )

                tab_preview, tab_edit, tab_diff, tab_json = st.tabs(["👁 預覽","📄 編輯","⚔️ 比對 (Diff)","📦 JSON"])

                with tab_preview:
                    tab_scroll_container = st.container()
                    with tab_scroll_container:
                        st.markdown('<div class="tab-scroll-marker"></div>', unsafe_allow_html=True)

                        if is_image(ext):
                            st.image(str(path), use_container_width=True)
                        elif is_pdf(ext):
                            b64 = file_b64(path)
                            st.markdown(
                                f'<iframe src="data:application/pdf;base64,{b64}" '
                                f'width="100%" height="560" class="pdf-frame"></iframe>',
                                unsafe_allow_html=True,
                            )
                        elif is_docx(ext):
                            try:
                                from docx import Document
                                doc  = Document(str(path))
                                text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                                st.markdown(text or "_（文件內容為空）_")
                            except Exception as e:
                                st.error(f"無法解析 DOCX：{e}")
                        elif ext == "html":
                            html_src = path.read_text(encoding="utf-8", errors="replace")
                            st.iframe(html_src, height=560)
                        elif ext == "svg":
                            svg_src = path.read_text(encoding="utf-8", errors="replace")
                            wrapped_svg = """
                            <!DOCTYPE html>
                            <html>
                            <head>
                              <script src="https://cdn.jsdelivr.net/npm/svg-pan-zoom@3.6.1/dist/svg-pan-zoom.min.js"></script>
                              <style>
                                html, body {
                                  margin: 0;
                                  padding: 0;
                                  width: 100%;
                                  height: 100%;
                                  overflow: hidden;
                                  background-color: #1e1e1e;
                                }
                                #svg-container {
                                  width: 100%;
                                  height: 100%;
                                }
                                svg {
                                  width: 100%;
                                  height: 100%;
                                }
                              </style>
                            </head>
                            <body>
                              <div id="svg-container">
                                __SVG_CONTENT__
                              </div>
                              <script>
                                window.onload = function() {
                                  var svgElement = document.querySelector('#svg-container svg');
                                  if (svgElement) {
                                    svgElement.setAttribute('width', '100%');
                                    svgElement.setAttribute('height', '100%');
                                    svgPanZoom(svgElement, {
                                      zoomEnabled: true,
                                      controlIconsEnabled: true,
                                      fit: true,
                                      center: true,
                                      minZoom: 0.1,
                                      maxZoom: 10
                                    });
                                  }
                                };
                              </script>
                            </body>
                            </html>
                            """.replace("__SVG_CONTENT__", svg_src)
                            st.iframe(wrapped_svg, height=560)
                        elif ext in {"md","txt"}:
                            st.markdown(path.read_text(encoding="utf-8", errors="replace"))
                        elif ext == "json":
                            try:
                                st.json(json.loads(path.read_text(encoding="utf-8", errors="replace")))
                            except json.JSONDecodeError:
                                st.code(path.read_text(encoding="utf-8", errors="replace"), language="json")
                        else:
                            st.info("此檔案類型暫不支援預覽，請下載後開啟。")

                        st.download_button(
                            f"⬇ 下載 {active['name']}",
                            data=path.read_bytes(),
                            file_name=active["name"],
                            mime="application/octet-stream",
                            use_container_width=True,
                        )

                with tab_edit:
                    tab_scroll_container = st.container()
                    with tab_scroll_container:
                        st.markdown('<div class="tab-scroll-marker"></div>', unsafe_allow_html=True)

                        if is_editable_text(ext):
                            saved_text = path.read_text(encoding="utf-8", errors="replace")

                            ace_lang = {
                                "html": "html",
                                "css": "css",
                                "js": "javascript",
                                "json": "json",
                                "md": "markdown",
                                "svg": "xml",
                            }.get(ext.lower(), "text")

                            edited_content = st_ace(
                                value=saved_text,
                                language=ace_lang,
                                theme="monokai",
                                height=420,
                                key=f"ace_editor_{active['id']}",
                            )

                            if edited_content != saved_text:
                                if st.button("💾 儲存並套用變更", key=f"save_edit_{active['id']}", type="primary", use_container_width=True):
                                    path.write_text(edited_content, encoding="utf-8")
                                    active["ts"] = datetime.now().strftime("%H:%M")
                                    st.success("變更已成功套用！")
                                    st.rerun()
                        else:
                            st.info("此類型目前不支援編輯。")

                with tab_diff:
                    tab_scroll_container = st.container()
                    with tab_scroll_container:
                        st.markdown('<div class="tab-scroll-marker"></div>', unsafe_allow_html=True)

                        if is_editable_text(ext):
                            orig = active.get("original_content", "")
                            curr = path.read_text(encoding="utf-8", errors="replace")
                            if not orig:
                                st.info("此檔案無原始版本記錄，無法進行比對。")
                            elif orig == curr:
                                st.success("✨ 目前內容與原始版本完全一致，無任何變更。")
                            else:
                                st.markdown("##### ⚔️ 原始版本 vs. 目前變更比對")
                                diff_html = generate_diff_html(orig, curr)
                                st.iframe(diff_html, height=520)
                        else:
                            st.info("此類型目前不支援版本比對。")

                with tab_json:
                    tab_scroll_container = st.container()
                    with tab_scroll_container:
                        st.markdown('<div class="tab-scroll-marker"></div>', unsafe_allow_html=True)

                        st.json({
                            "surface": {
                                "id":         active["id"],
                                "filename":   active["name"],
                                "ext":        active["ext"],
                                "source":     active["source"],
                                "ts":         active["ts"],
                                "path":       str(active["path"]),
                                "size_bytes": path.stat().st_size,
                            }
                        })
    else:
        # 展示欄折疊時顯示提示
        st.info("展示欄已折疊，點擊上方按鈕展開")
